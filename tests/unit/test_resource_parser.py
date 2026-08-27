from pathlib import Path
from unittest.mock import AsyncMock

import pymupdf as fitz
from docx import Document
from PIL import Image

from app.application.resources import ResourceService
from app.config import Settings


def test_markdown_parser_preserves_section_order(tmp_path: Path) -> None:
    path = tmp_path / "handout.md"
    path.write_text("# 稳定性\n内容A\n## 劳斯判据\n内容B", encoding="utf-8")

    sections = ResourceService._parse_markdown(path)

    assert [section.title for section in sections] == ["稳定性", "劳斯判据"]
    assert sections[0].sequence == 1
    assert "内容B" in sections[1].text


async def test_pdf_docx_and_image_parsers(tmp_path: Path) -> None:
    service = ResourceService(AsyncMock(), Settings(use_fake_model=True))

    pdf_path = tmp_path / "course.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Stability chapter")
    pdf.save(pdf_path)
    pdf.close()
    pdf_sections = await service._parse_file(pdf_path, "application/pdf")
    assert pdf_sections[0].method == "PYMUPDF"

    docx_path = tmp_path / "handout.docx"
    document = Document()
    document.add_heading("Root locus", level=1)
    document.add_paragraph("Exercise outline")
    document.save(docx_path)
    docx_sections = await service._parse_file(
        docx_path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert docx_sections[0].title == "Root locus"

    image_path = tmp_path / "scan.png"
    Image.new("RGB", (120, 80), "white").save(image_path)
    image_sections = await service._parse_file(image_path, "image/png")
    assert image_sections[0].method == "QWEN_VISION"
